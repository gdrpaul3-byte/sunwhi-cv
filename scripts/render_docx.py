#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Optional .docx exporters.

The Korean 이력서 is the one that really needs Word: Korean institutions expect
the boxed-table form. The English CV is exported too so it can be edited by
anyone who does not want to touch YAML.

Requires python-docx (`pip install python-docx`).
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from cvdata import (PUB_KINDS, author_string, citation, daterange, funded_grants,
                    grants, stamp)

KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"


def _base(doc: Document, font: str) -> None:
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    from docx.oxml.ns import qn
    rfonts.set(qn("w:eastAsia"), font)
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(50)
        section.top_margin = section.bottom_margin = Pt(45)


def _heading(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def _grid(doc: Document, headers: list[str], rows: list[list[str]],
          widths: list[int] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            run.font.size = Pt(9)
    return t


# --------------------------------------------------------------------------

def korean_resume(d: dict, out_dir: str) -> str:
    p = d["person"]
    doc = Document()
    _base(doc, KR_FONT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("이  력  서")
    r.bold = True
    r.font.size = Pt(20)

    _heading(doc, "1. 인적사항")
    rows = [
        ("성명(한글)", p.get("name_kr", "")),
        ("성명(영문)", "%s%s" % (p.get("name_en", ""),
                              " (%s)" % ", ".join(p["name_en_alt"]) if p.get("name_en_alt") else "")),
        ("생년월일", p.get("birth", "")),
        ("이메일", ", ".join(p.get("emails") or [])),
        ("핸드폰", (p.get("phones") or {}).get("mobile", "")),
        ("ORCID", p.get("orcid", "")),
        ("주소", p.get("address_kr", "")),
    ]
    _grid(doc, ["항목", "내용"], [[k, v] for k, v in rows if v])

    _heading(doc, "2. 학력사항")
    _grid(doc, ["기간", "학교명", "학위", "전공(학과)", "지도교수", "학위논문"],
          [[daterange(e.get("start"), e.get("end"), lang="kr"), e.get("org_kr", ""),
            e.get("degree_kr", ""), e.get("dept_kr", ""), e.get("advisor_kr", ""),
            e.get("thesis", "")] for e in d["education"]])

    _heading(doc, "3. 경력사항")
    _grid(doc, ["기간", "기관명", "부서", "직위"],
          [[daterange(a.get("start"), a.get("end"), lang="kr"),
            a.get("org_kr") or a.get("org_en", ""),
            a.get("dept_kr") or a.get("dept_en", ""),
            a.get("role_kr") or a.get("role_en", "")] for a in d["appointments"]])

    _heading(doc, "4. 논문 실적")
    pubrows = []
    for kind in PUB_KINDS:
        for pub in d["publications"][kind]:
            vol = str(pub.get("volume") or "")
            if pub.get("issue"):
                vol += "(%s)" % pub["issue"]
            if pub.get("article"):
                vol += (", " if vol else "") + str(pub["article"])
            pubrows.append([pub.get("year", ""), pub.get("venue", ""), vol,
                            pub.get("title", ""), pub.get("role_kr", ""),
                            pub.get("note_kr") or pub.get("note") or ""])
    _grid(doc, ["출간연도", "학술지", "권·호", "제목", "기여도", "특이사항"], pubrows)

    shown_grants = grants(d)
    if shown_grants:
        _heading(doc, "5. 연구비 실적")
        status_kr = {"awarded": "선정", "applied": "신청", "in_review": "심사중"}
        _grid(doc, ["기간", "과제명", "지원기관", "역할", "연구비", "상태"],
              [[g.get("period", ""), g.get("title_kr") or g.get("title_en", ""),
                g.get("funder_kr") or g.get("funder_en", ""),
                g.get("role_kr") or g.get("role", ""), g.get("amount", ""),
                status_kr.get(g.get("status"), "")] for g in shown_grants])

    _heading(doc, "6. 강의 실적")
    _grid(doc, ["연도·학기", "강의명", "교과코드", "학점", "학교명"],
          [[t.get("term", ""), t.get("name_kr") or t.get("name_en", ""),
            t.get("code", ""), t.get("credits", ""),
            t.get("institution_kr") or t.get("institution_en", "")]
           for t in d["teaching"]])

    if d["software"]:
        _heading(doc, "7. 프로그램 개발 실적")
        _grid(doc, ["연도", "프로그램명", "등록번호", "내용"],
              [[s.get("year", ""), s.get("name_kr") or s.get("name_en", ""),
                s.get("registration", ""),
                s.get("description_kr") or s.get("description_en", "")]
               for s in d["software"]])
        note = doc.add_paragraph()
        note.add_run("* 요청 시 등록 증명서 제공").font.size = Pt(8)

    if d["consulting"]:
        _heading(doc, "8. 외부 강연·워크샵·컨설팅")
        _grid(doc, ["기간", "기관", "형태", "주제"],
              [[c.get("period") or c.get("year", ""),
                c.get("org_kr") or c.get("org_en", ""),
                c.get("kind_kr") or c.get("kind_en", ""),
                c.get("topic_kr") or c.get("topic_en", "")]
               for c in d["consulting"]])

    if d["service"]:
        _heading(doc, "9. 교내·학회 활동")
        _grid(doc, ["기간", "활동", "기관"],
              [[s.get("period") or s.get("year", ""),
                s.get("role_kr") or s.get("role_en", ""), s.get("org", "")]
               for s in d["service"]])

    if d["certifications"]:
        _heading(doc, "10. 교육 이수 및 자격")
        _grid(doc, ["일자", "명칭", "발급기관"],
              [[c.get("date", ""), c.get("name_kr") or c.get("name_en", ""),
                c.get("issuer", "")] for c in d["certifications"]])

    if d["media"]:
        _heading(doc, "11. 언론 보도")
        _grid(doc, ["연도", "언론사", "기사 제목"],
              [[m.get("year", ""), m.get("outlet", ""), m.get("title", "")]
               for m in d["media"]])

    if d.get("military"):
        _heading(doc, "12. 병역사항")
        _grid(doc, ["항목", "내용"], [[k, v] for k, v in d["military"].items()])

    doc.add_paragraph()
    c = doc.add_paragraph("본 이력서에 기재한 사항은 사실과 다름없음을 확인합니다.")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp = doc.add_paragraph("작성일 : %s" % stamp(d))
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sp = doc.add_paragraph("지원자 : %s          (인)" % p.get("name_kr", ""))
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(out_dir, "이력서_%s.docx" % p.get("name_kr", "cv"))
    doc.save(path)
    return path


def english_cv(d: dict, out_dir: str) -> str:
    p = d["person"]
    doc = Document()
    _base(doc, EN_FONT)

    h = doc.add_paragraph()
    r = h.add_run("%s, %s" % (p.get("name_en"), p.get("degree", "Ph.D.")))
    r.bold = True
    r.font.size = Pt(18)
    if p.get("name_en_alt"):
        sub = doc.add_paragraph()
        sr = sub.add_run("Also published as %s" % ", ".join(p["name_en_alt"]))
        sr.italic = True
        sr.font.size = Pt(9)
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    info = doc.add_paragraph()
    info.add_run("%s, %s, %s\n" % (p.get("title_en"), p.get("dept_en"),
                                   p.get("institution_en"))).font.size = Pt(10)
    info.add_run("%s\n" % " · ".join(p.get("emails") or [])).font.size = Pt(9)
    ids = p.get("ids") or {}
    info.add_run("ORCID %s%s" % (
        p.get("orcid", ""),
        " · Google Scholar: scholar.google.com/citations?user=%s" % ids["scholar"]
        if ids.get("scholar") else "")).font.size = Pt(9)

    _heading(doc, "PROFESSIONAL APPOINTMENTS")
    for a in d["appointments"]:
        para = doc.add_paragraph()
        para.add_run("%s  " % daterange(a.get("start"), a.get("end"))).bold = True
        para.add_run("%s, %s%s" % (a.get("role_en", ""),
                                   "%s, " % a["dept_en"] if a.get("dept_en") else "",
                                   a.get("org_en", "")))
        if a.get("advisor"):
            para.add_run("\n    Advisor: %s" % a["advisor"]).italic = True

    _heading(doc, "EDUCATION")
    for e in d["education"]:
        para = doc.add_paragraph()
        para.add_run("%s  " % daterange(e.get("start"), e.get("end"))).bold = True
        para.add_run("%s, %s, %s" % (e.get("degree_en", ""), e.get("dept_en", ""),
                                     e.get("org_en", "")))
        if e.get("thesis"):
            para.add_run("\n    Dissertation: %s" % e["thesis"]).italic = True

    labels = {"journal": "PEER-REVIEWED PUBLICATIONS",
              "preprint": "PREPRINTS & MANUSCRIPTS",
              "conference": "CONFERENCE PROCEEDINGS"}
    for kind in PUB_KINDS:
        pubs = d["publications"][kind]
        if not pubs:
            continue
        _heading(doc, labels[kind])
        for i, pub in enumerate(pubs, 1):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            para.add_run("%d. %s" % (i, citation(pub))).font.size = Pt(9.5)

    if d["teaching"]:
        _heading(doc, "TEACHING")
        _grid(doc, ["Term", "Course", "Code", "Cr.", "Institution"],
              [[t.get("term", ""),
                "%s%s" % (t.get("name_en") or t.get("name_kr", ""),
                          " (%s)" % t["name_kr"] if t.get("name_en") and t.get("name_kr") else ""),
                t.get("code", ""), t.get("credits", ""),
                t.get("institution_en") or t.get("institution_kr", "")]
               for t in d["teaching"]])

    for key, label, fmt in [
        ("grants", "RESEARCH FUNDING",
         lambda g: "%s — %s, %s%s" % (g.get("period", ""), g.get("title_en") or g.get("title_kr"),
                                      g.get("funder_en") or g.get("funder_kr", ""),
                                      " (%s)" % g["amount"] if g.get("amount") else "")),
        ("invited_talks", "INVITED TALKS",
         lambda t: "%s — “%s.” %s%s" % (t.get("year", ""), (t.get("title") or "").rstrip("."),
                                        t.get("venue", ""),
                                        ", %s" % t["location"] if t.get("location") else "")),
        ("posters", "CONFERENCE PRESENTATIONS",
         lambda t: "%s. “%s.” %s, %s." % (", ".join(t.get("authors") or []),
                                          (t.get("title") or "").rstrip("."),
                                          t.get("venue", ""), t.get("year", ""))),
        ("awards", "AWARDS & HONORS",
         lambda a: "%s — %s%s" % (a.get("year", ""), a.get("name_en") or a.get("name_kr"),
                                  ", %s" % a["org"] if a.get("org") else "")),
        ("service", "SERVICE",
         lambda s: "%s — %s%s" % (s.get("period") or s.get("year", ""),
                                  s.get("role_en") or s.get("role_kr"),
                                  ", %s" % s["org"] if s.get("org") else "")),
        ("consulting", "AI AGENT ONBOARDING & CONSULTING",
         lambda c: "%s — %s%s%s" % (c.get("period") or c.get("year", ""),
                                    c.get("org_en") or c.get("org_kr"),
                                    ", %s" % c["kind_en"] if c.get("kind_en") else "",
                                    ": %s" % c["topic_en"] if c.get("topic_en") else "")),
        ("mentoring", "MENTORING",
         lambda m: "%s — %s%s" % (m.get("period", ""), m.get("name", ""),
                                  ", %s" % m["role"] if m.get("role") else "")),
        ("certifications", "CERTIFICATIONS & TRAINING",
         lambda c: "%s — %s%s" % (c.get("date", ""), c.get("name_en") or c.get("name_kr"),
                                  ", %s" % c["issuer"] if c.get("issuer") else "")),
        ("media", "MEDIA COVERAGE",
         lambda m: "%s — %s, %s" % (m.get("year", ""), m.get("title", ""), m.get("outlet", ""))),
    ]:
        items = d.get(key) or []
        if key == "grants":
            items = funded_grants(d)
        if not items:
            continue
        _heading(doc, label)
        for item in items:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            para.add_run(fmt(item)).font.size = Pt(9.5)

    if d["peer_review"]:
        _heading(doc, "PEER REVIEW")
        for r_ in d["peer_review"]:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run("%s (%s)" % (r_.get("journal", ""), r_.get("years", ""))).font.size = Pt(9.5)

    if d["references"]:
        _heading(doc, "REFERENCES")
        for ref in d["references"]:
            para = doc.add_paragraph()
            para.add_run("%s — %s, %s (%s)" % (ref.get("name", ""), ref.get("title", ""),
                                               ref.get("org", ""), ref.get("email", ""))).font.size = Pt(9.5)

    path = os.path.join(out_dir, "CV_%s.docx" % (p.get("name_en") or "cv").replace(" ", ""))
    doc.save(path)
    return path


def build(d: dict, out_dir: str) -> list[str]:
    return [english_cv(d, out_dir), korean_resume(d, out_dir)]
